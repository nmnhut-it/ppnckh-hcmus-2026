#!/usr/bin/env python3
"""
Convert markdown thesis proposal to docx matching the reference document style.

Reference: _De cuong luan van thac si_CH1501015_Le Quang Thai.doc
- Page: A4 (21.0 x 29.7 cm)
- Margins: Left/Right ~3.17 cm, Top/Bottom ~2.01 cm
- Primary font: Times New Roman
- Body text: 13pt, 1.5 line spacing
- Heading styles matching Vietnamese academic standard
"""

import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Configuration matching reference doc ──────────────────────────────────────

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(13)
FONT_SIZE_HEADING1 = Pt(14)
FONT_SIZE_HEADING2 = Pt(13)
FONT_SIZE_HEADING3 = Pt(13)
FONT_SIZE_HEADING4 = Pt(13)
FONT_SIZE_TABLE = Pt(12)
FONT_SIZE_REFERENCE = Pt(12)

LINE_SPACING = 1.5
SPACE_AFTER_PARA = Pt(6)
SPACE_BEFORE_HEADING = Pt(12)

PAGE_WIDTH = Cm(21.0)
PAGE_HEIGHT = Cm(29.7)
MARGIN_LEFT = Cm(3.17)
MARGIN_RIGHT = Cm(3.17)
MARGIN_TOP = Cm(2.01)
MARGIN_BOTTOM = Cm(2.01)
HEADER_DISTANCE = Cm(1.27)
FOOTER_DISTANCE = Cm(1.27)


def setup_document():
    """Create and configure a new document with page setup."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE
    section.orientation = WD_ORIENT.PORTRAIT

    # Set default font for the document
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_BODY
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.space_after = SPACE_AFTER_PARA

    # Set East Asian font
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), FONT_NAME)

    return doc


def set_run_font(run, font_name=FONT_NAME, font_size=None, bold=None, italic=None,
                 color=None, underline=None):
    """Configure a run's font properties."""
    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = color
    if underline is not None:
        run.font.underline = underline

    # Set East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), font_name)


def add_formatted_text(paragraph, text, font_size=FONT_SIZE_BODY, base_bold=False,
                       base_italic=False):
    """Parse inline markdown formatting and add runs to a paragraph."""
    # Process inline formatting: **bold**, *italic*, ***bold italic***
    # Also handle [text](url) links
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'   # ***bold italic***
        r'|(\*\*(.+?)\*\*)'      # **bold**
        r'|(\*(.+?)\*)'          # *italic*
        r'|(\[([^\]]+)\]\(([^)]+)\))'  # [text](url)
    )

    last_end = 0
    for m in pattern.finditer(text):
        # Add text before this match
        if m.start() > last_end:
            plain = text[last_end:m.start()]
            if plain:
                run = paragraph.add_run(plain)
                set_run_font(run, font_size=font_size, bold=base_bold,
                             italic=base_italic)

        if m.group(2):  # ***bold italic***
            run = paragraph.add_run(m.group(2))
            set_run_font(run, font_size=font_size, bold=True, italic=True)
        elif m.group(4):  # **bold**
            run = paragraph.add_run(m.group(4))
            set_run_font(run, font_size=font_size, bold=True, italic=base_italic)
        elif m.group(6):  # *italic*
            run = paragraph.add_run(m.group(6))
            set_run_font(run, font_size=font_size, bold=base_bold, italic=True)
        elif m.group(8):  # [text](url)
            link_text = m.group(9)
            run = paragraph.add_run(link_text)
            set_run_font(run, font_size=font_size, bold=base_bold,
                         italic=base_italic, color=RGBColor(0x00, 0x00, 0xFF),
                         underline=True)

        last_end = m.end()

    # Add remaining text
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            run = paragraph.add_run(remaining)
            set_run_font(run, font_size=font_size, bold=base_bold,
                         italic=base_italic)


def add_heading_paragraph(doc, text, level, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a heading paragraph with proper formatting."""
    para = doc.add_paragraph()
    para.alignment = alignment

    if level == 1:
        font_size = FONT_SIZE_HEADING1
        bold = True
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
    elif level == 2:
        font_size = FONT_SIZE_HEADING2
        bold = True
        para.paragraph_format.space_before = SPACE_BEFORE_HEADING
        para.paragraph_format.space_after = SPACE_AFTER_PARA
    elif level == 3:
        font_size = FONT_SIZE_HEADING3
        bold = True
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = SPACE_AFTER_PARA
    else:
        font_size = FONT_SIZE_HEADING4
        bold = True
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = SPACE_AFTER_PARA

    para.paragraph_format.line_spacing = LINE_SPACING

    add_formatted_text(para, text, font_size=font_size, base_bold=bold)
    return para


def add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_text(para, header.strip(), font_size=FONT_SIZE_TABLE,
                           base_bold=True)
        # Shade header
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= num_cols:
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            add_formatted_text(para, cell_text.strip(), font_size=FONT_SIZE_TABLE)

    # Set table width to fill page
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn("w:type"), "pct")
        tblW.set(qn("w:w"), "5000")

    doc.add_paragraph()  # spacing after table
    return table


def parse_table_block(lines):
    """Parse a markdown table block into headers and rows."""
    if len(lines) < 2:
        return None, None

    def split_row(line):
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        return [c.strip() for c in line.split("|")]

    headers = split_row(lines[0])

    # Skip separator line (line with dashes)
    rows = []
    for line in lines[2:]:
        if line.strip():
            rows.append(split_row(line))

    return headers, rows


def process_markdown(md_text):
    """Convert markdown text to a formatted docx document."""
    doc = setup_document()
    lines = md_text.split("\n")
    i = 0
    in_reference_section = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---" or stripped == "***" or stripped == "___":
            # Add a thin horizontal line
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            pPr = para._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)

            # Check if this is a reference section
            if "Tài liệu tham khảo" in text:
                in_reference_section = True

            # Title-level headings (level 1) are centered
            if level == 1:
                add_heading_paragraph(doc, text, level,
                                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                add_heading_paragraph(doc, text, level)
            i += 1
            continue

        # Table detection
        if "|" in stripped and i + 1 < len(lines) and re.match(
            r'^\|?\s*[-:]+[-|\s:]*$', lines[i + 1].strip()
        ):
            table_lines = []
            j = i
            while j < len(lines) and "|" in lines[j].strip():
                table_lines.append(lines[j])
                j += 1
            headers, rows = parse_table_block(table_lines)
            if headers and rows:
                add_table(doc, headers, rows)
            i = j
            continue

        # Bullet list items
        bullet_match = re.match(r'^(\s*)-\s+(.+)$', stripped)
        if bullet_match:
            text = bullet_match.group(2)
            para = doc.add_paragraph()
            para.style = doc.styles["Normal"]
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.first_line_indent = Cm(-0.5)
            para.paragraph_format.line_spacing = LINE_SPACING
            para.paragraph_format.space_after = Pt(3)

            # Add bullet character
            run = para.add_run("- ")
            set_run_font(run, font_size=FONT_SIZE_BODY)
            add_formatted_text(para, text, font_size=FONT_SIZE_BODY)

            # Check for continuation lines (indented lines that follow)
            while (i + 1 < len(lines) and lines[i + 1].strip()
                   and not lines[i + 1].strip().startswith("-")
                   and not lines[i + 1].strip().startswith("#")
                   and not lines[i + 1].strip().startswith("|")
                   and not re.match(r'^\d+\.', lines[i + 1].strip())
                   and not lines[i + 1].strip().startswith(">")):
                i += 1
                cont_text = lines[i].strip()
                if cont_text:
                    run = para.add_run(" ")
                    set_run_font(run, font_size=FONT_SIZE_BODY)
                    add_formatted_text(para, cont_text, font_size=FONT_SIZE_BODY)

            i += 1
            continue

        # Sub-bullet list items (indented with spaces)
        sub_bullet_match = re.match(r'^(\s+)-\s+(.+)$', line)
        if sub_bullet_match and len(sub_bullet_match.group(1)) >= 2:
            text = sub_bullet_match.group(2)
            para = doc.add_paragraph()
            para.style = doc.styles["Normal"]
            para.paragraph_format.left_indent = Cm(1.5)
            para.paragraph_format.first_line_indent = Cm(-0.5)
            para.paragraph_format.line_spacing = LINE_SPACING
            para.paragraph_format.space_after = Pt(3)

            run = para.add_run("- ")
            set_run_font(run, font_size=FONT_SIZE_BODY)
            add_formatted_text(para, text, font_size=FONT_SIZE_BODY)
            i += 1
            continue

        # Numbered list items
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if numbered_match:
            num = numbered_match.group(1)
            text = numbered_match.group(2)
            para = doc.add_paragraph()
            para.style = doc.styles["Normal"]
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.first_line_indent = Cm(-0.5)
            para.paragraph_format.line_spacing = LINE_SPACING
            para.paragraph_format.space_after = Pt(3)

            run = para.add_run(f"{num}. ")
            set_run_font(run, font_size=FONT_SIZE_BODY)
            add_formatted_text(para, text, font_size=FONT_SIZE_BODY)
            i += 1
            continue

        # Reference entries [N]
        ref_match = re.match(r'^\[(\d+)\]\s+(.+)$', stripped)
        if ref_match:
            num = ref_match.group(1)
            text = ref_match.group(2)
            para = doc.add_paragraph()
            para.style = doc.styles["Normal"]
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.first_line_indent = Cm(-1.0)
            para.paragraph_format.line_spacing = LINE_SPACING
            para.paragraph_format.space_after = Pt(2)

            run = para.add_run(f"[{num}] ")
            set_run_font(run, font_size=FONT_SIZE_REFERENCE)
            add_formatted_text(para, text, font_size=FONT_SIZE_REFERENCE)

            # Handle continuation lines
            while (i + 1 < len(lines) and lines[i + 1].strip()
                   and not re.match(r'^\[(\d+)\]', lines[i + 1].strip())
                   and not lines[i + 1].strip().startswith("#")
                   and not lines[i + 1].strip().startswith("---")
                   and not lines[i + 1].strip().startswith("|")):
                i += 1
                cont = lines[i].strip()
                if cont:
                    run = para.add_run(" ")
                    set_run_font(run, font_size=FONT_SIZE_REFERENCE)
                    add_formatted_text(para, cont, font_size=FONT_SIZE_REFERENCE)

            i += 1
            continue

        # Signature table (last table with Giảng viên / Học viên)
        if "Giảng viên hướng dẫn" in stripped and "|" in stripped:
            # Parse as a simple 2-column signature block
            table_lines = []
            j = i
            while j < len(lines) and "|" in lines[j].strip():
                table_lines.append(lines[j])
                j += 1

            if len(table_lines) >= 2:
                # Create a borderless table for signatures
                sig_table = doc.add_table(rows=2, cols=2)
                sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Remove borders
                for row in sig_table.rows:
                    for cell in row.cells:
                        tcPr = cell._element.get_or_add_tcPr()
                        tcBorders = parse_xml(
                            f'<w:tcBorders {nsdecls("w")}>'
                            f'<w:top w:val="none" w:sz="0" w:space="0"/>'
                            f'<w:left w:val="none" w:sz="0" w:space="0"/>'
                            f'<w:bottom w:val="none" w:sz="0" w:space="0"/>'
                            f'<w:right w:val="none" w:sz="0" w:space="0"/>'
                            f'</w:tcBorders>'
                        )
                        tcPr.append(tcBorders)

                # Header row
                for idx, header_text in enumerate(
                    ["Giảng viên hướng dẫn", "Học viên cao học"]
                ):
                    cell = sig_table.rows[0].cells[idx]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(header_text)
                    set_run_font(run, font_size=FONT_SIZE_BODY, bold=True)

                # Signature row
                for idx, sig_text in enumerate(["(Ký tên)", "(Ký tên)"]):
                    cell = sig_table.rows[1].cells[idx]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(sig_text)
                    set_run_font(run, font_size=FONT_SIZE_BODY, italic=True)

            i = j
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        para.style = doc.styles["Normal"]
        para.paragraph_format.line_spacing = LINE_SPACING
        para.paragraph_format.space_after = SPACE_AFTER_PARA
        # First line indent for body paragraphs (Vietnamese academic style)
        if not stripped.startswith("*") or stripped.startswith("**"):
            para.paragraph_format.first_line_indent = Cm(1.0)

        # Handle lines that start with only italic (like the English title)
        if stripped.startswith("*(") and stripped.endswith(")*"):
            para.paragraph_format.first_line_indent = None
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inner_text = stripped[1:-1]  # Remove outer * *
            run = para.add_run(inner_text)
            set_run_font(run, font_size=FONT_SIZE_BODY, italic=True)
            i += 1
            continue

        add_formatted_text(para, stripped, font_size=FONT_SIZE_BODY)
        i += 1

    return doc


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, "de-cuong-luan-van-voice-assessment.md")
    docx_path = os.path.join(
        script_dir, "de-cuong-luan-van-voice-assessment.docx"
    )

    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = process_markdown(md_text)
    doc.save(docx_path)
    print(f"Converted: {md_path}")
    print(f"Output:    {docx_path}")


if __name__ == "__main__":
    main()
